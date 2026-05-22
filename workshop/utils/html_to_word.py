"""Utilitaire pour générer des documents Word."""

import json

from docx import Document


def create_word_from_stats(context_data):
    """
    Créer un document Word à partir des données de statistiques d'ateliers.

    Args:
        context_data (dict): Dictionnaire contenant les données de stats.

    Returns:
        Document: Le document Word généré avec les statistiques formatées.
    """
    # Créer un nouveau document Word
    doc = Document()

    # Ajouter le titre principal
    doc.add_heading(
        f'Statistiques des ateliers {context_data.get("current_year", "")}', 0
    )

    # Ajouter un avertissement
    doc.add_paragraph(
        "Ces statistiques incluent tous les ateliers de l'année"
        f" {context_data.get('current_year', '')}."
    )

    # Statistiques de participation
    doc.add_heading("Statistiques de participation", level=1)
    if "participation_stats" in context_data:
        stats = context_data["participation_stats"]
        p = doc.add_paragraph()
        p.add_run("Moyenne de participants : ").bold = True
        p.add_run(f"{stats.get('moyenne', 0):.1f}")

        p = doc.add_paragraph()
        p.add_run("Total des participants : ").bold = True
        p.add_run(str(stats.get("total_participants", 0)))

        p = doc.add_paragraph()
        p.add_run("Nombre total d'ateliers : ").bold = True
        p.add_run(str(stats.get("total_ateliers", 0)))

        p = doc.add_paragraph()
        p.add_run("Nombre d'accueils de classe : ").bold = True
        p.add_run(str(stats.get("total_accueil_classe", 0)))

        p = doc.add_paragraph()
        p.add_run("Nombre d'ateliers (hors accueils de classe) : ").bold = True
        p.add_run(str(stats.get("total_hors_accueil_classe", 0)))

    # Statistiques mensuelles
    doc.add_heading("Évolution mensuelle", level=1)
    if "workshops_by_month" in context_data:
        months = [
            "Janvier",
            "Février",
            "Mars",
            "Avril",
            "Mai",
            "Juin",
            "Juillet",
            "Août",
            "Septembre",
            "Octobre",
            "Novembre",
            "Décembre",
        ]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Mois"
        header_cells[1].text = "Nombre d'ateliers"
        header_cells[2].text = "Nombre d'inscrits"

        # Parse les données JSON si nécessaire
        monthly_data = context_data["workshops_by_month"]
        if isinstance(monthly_data, str):
            monthly_data = json.loads(monthly_data)

        for month_data in monthly_data:
            row_cells = table.add_row().cells
            row_cells[0].text = months[month_data["month"] - 1]
            row_cells[1].text = str(month_data["count"])
            row_cells[2].text = str(month_data["total_registered"])

    # Canaux de communication
    doc.add_heading("Canaux de communication", level=1)
    if "communication_stats" in context_data:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Canal"
        header_cells[1].text = "Nombre d'ateliers"

        channels = {
            "instagram": "Instagram",
            "facebook": "Facebook",
            "mail": "Email",
            "portail": "Portail",
            "vdn": "VDN",
        }

        for key, label in channels.items():
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(
                context_data["communication_stats"].get(key, 0)
            )

    return doc


def save_stats_to_word(context_data, output_file):
    """
    Créer et sauvegarder un document Word.

    Args:
        context_data (dict): Dictionnaire contenant les données.
        output_file (str): Nom du fichier Word à créer.

    Returns:
        Document: Le document Word créé.
    """
    doc = create_word_from_stats(context_data)
    return doc
